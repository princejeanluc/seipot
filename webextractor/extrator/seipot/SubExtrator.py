from abc import ABC , abstractmethod
import pdfplumber
import docx
from odf import text , teletype
from odf.opendocument import load

class SubExtractor(ABC):
    def __init__(self):
        pass

    #renvoi true si elle prend en charge le type de fichier
    @abstractmethod
    def recognise(self,extension: str):
        pass

    def cleanText(self, list_text :str):
        for i in range(len(list_text)):
            list_text[i] = list_text[i].replace("\n",". ")
            list_text[i] = list_text[i].replace("·",",")
            list_text[i] = list_text[i].replace("~"," ")
            list_text[i] = list_text[i].replace("•",",")
            list_text[i] = " ".join(list_text[i].split())
        return list_text

    def extract(self,file_path:str):
        list_text = self.extractText(file_path)
        return self.cleanText(list_text)

    #retourne le texte contenu dans le document
    @abstractmethod
    def extractText(self, file_path:str):
        pass


class PDFTextExtractor (SubExtractor):
    def __init__(self):
        pass

    def recognise(self,extension: str):
        return extension == "pdf"

    def extractText(self, file_path:str):
        list_text = []
        with pdfplumber.open(file_path) as doc:
            for page in doc.pages :
                list_text.append(page.extract_text())
        return list_text




class DocTextExtractor(SubExtractor):
    def __init__(self):
        pass

    def recognise(self,extension: str):
        return extension =="docx"

    def extractText(self, file_path:str):
        list_text = []
        document = docx.Document(file_path)
        content = ""
        for i in range(len(document.paragraphs)):
            content+= document.paragraphs[i].text + " "
            if i == len(document.paragraphs)-1 or (i % 30 == 0 and  i!=0) :
                list_text.append(content)
                content = ""
        return list_text


class OdtTextExtrator(SubExtractor):
    def __init__(self):
        pass

    def recognise(self,extension: str):
        return extension =="odt"

    def extractText(self, file_path:str):
        content = ""
        document = load(file_path)
        list_text =[]
        for i in range(len(document.getElementsByType(text.P))):
            content+= teletype.extractText(document.getElementsByType(text.P)[i]) + " "
            if i == len(document.getElementsByType(text.P))-1 or (i % 30 == 0 and  i!=0) :
                list_text.append(content)
                content = ""
        return list_text

class FileTextExtractor(SubExtractor):
    def __init__(self):
        pass

    def recognise(self,extension: str):
        return extension =="txt"

    def extractText(self, file_path:str):
        content = ""
        with open(file_path, "r", encoding="utf-8",errors="ignore") as f:
            content = f.read()
        return [content]

if __name__ == "__main__":
    pass
    #print(OdtTextExtrator().extractText("./document.odt"))